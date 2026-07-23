from pathlib import Path
from unittest.mock import MagicMock, patch

from app.services.parsers.docx import parse_docx
from app.services.parsers.pdf import parse_pdf
from app.services.parsers.pptx import parse_pptx
from app.services.parsers.xlsx import parse_xlsx

FX = Path(__file__).parent / "fixtures"


def test_docx_section_path_from_headings(tmp_path):
    from docx import Document as D

    d = D()
    d.add_heading("第一章 总则", level=1)
    d.add_paragraph("本章规定了适用范围与基本原则。" * 200)  # >2000 chars -> forces a flush
    d.add_heading("1.1 定义", level=2)
    d.add_paragraph("本合同中的术语含义如下。" * 200)  # another section, now under L1 > L2
    f = tmp_path / "h.docx"
    d.save(f)

    units = parse_docx(f)
    assert units
    assert all(isinstance(u["section_path"], list) for u in units)
    paths = [u["section_path"] for u in units]
    assert ["第一章 总则"] in paths                    # section under level-1 only
    assert ["第一章 总则", "1.1 定义"] in paths          # breadcrumb under level-1 > level-2


def test_xlsx_section_path_is_sheet_name():
    units = parse_xlsx(FX / "tiny.xlsx")
    for u in units:
        assert len(u["section_path"]) == 1
        assert u["section_path"][0] in u["text"]


def test_pptx_section_path_is_list():
    units = parse_pptx(FX / "tiny.pptx")
    assert units
    assert all(isinstance(u["section_path"], list) for u in units)
    # tiny.pptx slides carry titles -> at least one section_path is a non-empty single title
    assert any(len(u["section_path"]) == 1 and u["section_path"][0].strip() for u in units)


@patch("app.services.parsers.pdf.PdfReader")
@patch("app.services.parsers.pdf.ocr_enabled", return_value=False)
def test_pdf_section_path_empty(mock_en, mock_reader):
    page = MagicMock()
    page.extract_text.return_value = "some native text " * 10
    mock_reader.return_value.pages = [page]
    pages = parse_pdf(FX / "scanned.pdf")
    assert pages and pages[0]["section_path"] == []
